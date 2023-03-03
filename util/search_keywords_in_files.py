import os
import re
import json
import docx
import openpyxl


def search_keywords_in_files(path, keyword, replace_str=None):
    result = []
    for root, dirs, files in os.walk(path):
        for filename in files:
            abs_path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            if ext in ('.doc', '.docx'):
                try:
                    doc = docx.Document(abs_path)
                    doc_content = ''
                    for para in doc.paragraphs:
                        doc_content += para.text
                    if keyword in doc_content:
                        count = doc_content.count(keyword)
                        item = {'name': filename, 'path': abs_path, 'count': count, 'status': 'read'}
                        if replace_str:
                            new_content = re.sub(keyword, replace_str, doc_content)
                            doc.paragraphs[0].text = new_content
                            doc.save(abs_path)
                            item['status'] = 'replaced'
                        result.append(item)
                except Exception as e:
                    print(f"Error reading {abs_path}: {e}")

            elif ext == '.xlsx' or ext == '.xls':
                try:
                    wb = openpyxl.load_workbook(abs_path)
                    cell_content = ''
                    for sheet in wb:
                        for row in sheet.iter_rows():
                            for cell in row:
                                cell_content += str(cell.value) + ' '
                    if keyword in cell_content:
                        count = cell_content.count(keyword)
                        item = {'name': filename, 'path': abs_path, 'count': count, 'status': 'read'}
                        if replace_str:
                            new_content = re.sub(keyword, replace_str, cell_content)
                            for sheet in wb:
                                for row in sheet.iter_rows():
                                    for cell in row:
                                        cell.value = new_content
                            wb.save(abs_path)
                            item['status'] = 'replaced'
                        result.append(item)
                except Exception as e:
                    print(f"Error reading {abs_path}: {e}")

            elif ext == '.json':
                try:
                    with open(abs_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    json_str = json.dumps(data)
                    if keyword in json_str:
                        count = json_str.count(keyword)
                        item = {'name': filename, 'path': abs_path, 'count': count, 'status': 'read'}
                        if replace_str:
                            new_content = json_str.replace(keyword, replace_str)
                            new_data = json.loads(new_content)
                            with open(abs_path, 'w', encoding='utf-8') as f:
                                json.dump(new_data, f, ensure_ascii=False, indent=4)
                            item['status'] = 'replaced'
                        result.append(item)
                except Exception as e:
                    print(f"Error reading {abs_path}: {e}")

            elif ext == '.txt' or ext == '.md' or ext == '.js' or ext == '.coffee' or ext == '.css' or ext == '.scss' or ext == '.sass':
                try:
                    with open(abs_path, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                    if keyword in file_content:
                        count = file_content.count(keyword)
                        item = {'name': filename, 'path': abs_path, 'count': count, 'status': 'read'}
                        if replace_str:
                            new_content = file_content.replace(keyword, replace_str)
                            with open(abs_path, 'w', encoding='utf-8') as f:
                                f.write(new_content)
                            item['status'] = 'replaced'
                        result.append(item)
                except Exception as e:
                    print(f"Error reading {abs_path}: {e}")
            else:
                item = {'name': filename, 'path': abs_path, 'count': 0, 'status': 'no support'}
                result.append(item)

    result = sorted(result, key=lambda x: x['count'], reverse=True)

    return result


ret = search_keywords_in_files('/opt/node/_trash', 'if')

for it in ret:
    if it['status'] != 'no support':
        print(f"{it['name']}, {it['count']} \n{it['path']}\n")
